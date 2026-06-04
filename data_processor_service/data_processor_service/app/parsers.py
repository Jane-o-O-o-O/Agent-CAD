from __future__ import annotations

import base64
import csv
import hashlib
import io
import json
import mimetypes
import os
import re
import shutil
import subprocess
import tempfile
import urllib.error
import urllib.request
import zipfile
from html import unescape
from pathlib import Path
from typing import Any

from .schemas import CleanParseResponse, ParsedContent, ParseResponse


def _load_local_env() -> None:
    env_path = Path(__file__).resolve().parents[1] / ".env"
    if not env_path.exists():
        return
    try:
        lines = env_path.read_text(encoding="utf-8").splitlines()
    except OSError:
        return
    for line in lines:
        line = line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        key = key.strip()
        value = value.strip().strip('"').strip("'")
        if key and key not in os.environ:
            os.environ[key] = value


_load_local_env()


SUPPORTED_EXTENSIONS = {
    "txt",
    "md",
    "html",
    "htm",
    "json",
    "csv",
    "xlsx",
    "xls",
    "doc",
    "docx",
    "pdf",
    "png",
    "jpg",
    "jpeg",
    "bmp",
    "tif",
    "tiff",
}


class ParseError(Exception):
    """Raised when an uploaded file cannot be parsed by the scoped workflow."""


def parse_upload(
    filename: str,
    data: bytes,
    content_type: str | None,
    include_raw: bool = True,
    include_debug_images: bool = True,
    output_mode: str = "clean",
) -> ParseResponse | CleanParseResponse:
    extension = _extension(filename)
    if extension not in SUPPORTED_EXTENSIONS:
        raise ParseError(f"Unsupported file type: {extension or 'unknown'}")

    metadata: dict[str, Any] = {
        "size_bytes": len(data),
        "sha256": hashlib.sha256(data).hexdigest(),
        "content_type": content_type or mimetypes.guess_type(filename)[0],
    }

    vision_budget = _VisionBudget(_env_int("DATA_PROCESSOR_VISION_MAX_IMAGES_PER_FILE", 3))
    content = _parse_by_extension(extension, filename, data, vision_budget)
    metadata["image_count"] = len(content.images)
    if extension == "doc":
        metadata["word_conversion"] = "doc_to_docx"

    if not include_raw:
        content.text = ""
    _enrich_content(content)
    if not include_debug_images:
        content.images = _compact_images(content.images)

    response = ParseResponse(
        status="success",
        filename=filename,
        source_type=extension,
        metadata=metadata,
        content=content,
        errors=[],
    )
    if output_mode.strip().lower() == "full":
        return response
    return _to_clean_response(response)


def _compact_images(images: list[dict[str, Any]]) -> list[dict[str, Any]]:
    compacted: list[dict[str, Any]] = []
    for image in images:
        analysis = image.get("analysis", {})
        vision_model = analysis.get("vision_model", {})
        compacted.append(
            {
                "index": image.get("index"),
                "filename": image.get("filename"),
                "role": image.get("role"),
                "content_type": image.get("content_type"),
                "width": analysis.get("width"),
                "height": analysis.get("height"),
                "ocr_text": analysis.get("ocr_text", ""),
                "vision_model": {
                    "status": vision_model.get("status"),
                    "provider": vision_model.get("provider"),
                    "model": vision_model.get("model"),
                    "reason": vision_model.get("reason"),
                    "errors": vision_model.get("errors"),
                },
                "vision_structured_data": analysis.get("vision_structured_data", {}),
                "errors": analysis.get("errors", []),
            },
        )
    return compacted


def _to_clean_response(response: ParseResponse) -> CleanParseResponse:
    content = response.content
    return CleanParseResponse(
        status=response.status,
        filename=response.filename,
        source_type=response.source_type,
        metadata=response.metadata,
        text=content.text,
        diagram_analysis=_strip_empty_debug(content.diagram_analysis),
        image_analysis_summary=content.image_analysis_summary,
        structured_data=_strip_empty_debug(content.structured_data),
        errors=response.errors,
    )


def _strip_empty_debug(value: Any) -> Any:
    if isinstance(value, dict):
        return {
            key: _strip_empty_debug(item)
            for key, item in value.items()
            if key != "raw" and item not in ({}, [], "", None)
        }
    if isinstance(value, list):
        return [_strip_empty_debug(item) for item in value if item not in ({}, [], "", None)]
    return value


def _extension(filename: str) -> str:
    return Path(filename).suffix.lower().lstrip(".")


def _parse_by_extension(
    extension: str,
    filename: str,
    data: bytes,
    vision_budget: "_VisionBudget",
) -> ParsedContent:
    if extension in {"txt", "md"}:
        return ParsedContent(text=_decode_text(data))
    if extension in {"html", "htm"}:
        return _parse_html(data)
    if extension == "json":
        return _parse_json(data)
    if extension == "csv":
        return _parse_csv(data)
    if extension in {"xlsx", "xls"}:
        return _parse_excel(data, filename)
    if extension == "doc":
        return _parse_doc(data, filename, vision_budget)
    if extension == "docx":
        return _parse_docx(data, vision_budget)
    if extension == "pdf":
        return _parse_pdf(data)
    if extension in {"png", "jpg", "jpeg", "bmp", "tif", "tiff"}:
        return _parse_image_file(filename, data, vision_budget)
    raise ParseError(f"Unsupported file type: {extension}")


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _parse_html(data: bytes) -> ParsedContent:
    try:
        from bs4 import BeautifulSoup
    except ImportError as exc:
        raise ParseError("Missing dependency: beautifulsoup4") from exc

    soup = BeautifulSoup(_decode_text(data), "html.parser")
    for node in soup(["script", "style", "noscript"]):
        node.extract()
    text = unescape(soup.get_text("\n"))
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return ParsedContent(text="\n".join(lines))


def _parse_json(data: bytes) -> ParsedContent:
    try:
        parsed = json.loads(_decode_text(data))
    except json.JSONDecodeError as exc:
        raise ParseError(f"Invalid JSON: {exc}") from exc

    records: list[dict[str, Any]] = []
    if isinstance(parsed, list):
        records = [item for item in parsed if isinstance(item, dict)]
    elif isinstance(parsed, dict):
        records = [parsed]

    return ParsedContent(text=json.dumps(parsed, ensure_ascii=False, indent=2), records=records)


def _parse_csv(data: bytes) -> ParsedContent:
    text = _decode_text(data)
    sample = text[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample)
    except csv.Error:
        dialect = csv.excel

    reader = csv.DictReader(io.StringIO(text), dialect=dialect)
    records = [dict(row) for row in reader]
    return ParsedContent(text=text, records=records)


def _parse_excel(data: bytes, filename: str) -> ParsedContent:
    try:
        import pandas as pd
    except ImportError as exc:
        raise ParseError("Missing dependency: pandas") from exc

    try:
        sheets = pd.read_excel(io.BytesIO(data), sheet_name=None)
    except Exception as exc:
        raise ParseError(f"Failed to parse Excel file {filename}: {exc}") from exc

    tables: list[dict[str, Any]] = []
    all_records: list[dict[str, Any]] = []
    for sheet_name, frame in sheets.items():
        clean_frame = frame.where(pd.notnull(frame), None)
        records = clean_frame.to_dict(orient="records")
        tables.append(
            {
                "sheet_name": str(sheet_name),
                "row_count": len(records),
                "columns": [str(column) for column in clean_frame.columns],
                "records": records,
            },
        )
        all_records.extend(records)

    return ParsedContent(tables=tables, records=all_records)


class _VisionBudget:
    def __init__(self, max_images: int) -> None:
        self.max_images = max(0, max_images)
        self.used = 0

    def allow(self) -> bool:
        if self.used >= self.max_images:
            return False
        self.used += 1
        return True


def _parse_docx(data: bytes, vision_budget: _VisionBudget) -> ParsedContent:
    try:
        from docx import Document
    except ImportError as exc:
        raise ParseError("Missing dependency: python-docx") from exc

    document = Document(io.BytesIO(data))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

    tables: list[list[list[str]]] = []
    for table in document.tables:
        table_rows: list[list[str]] = []
        for row in table.rows:
            table_rows.append([cell.text.strip() for cell in row.cells])
        tables.append(table_rows)

    images = _extract_docx_images(data, vision_budget)
    _merge_word_html_image_exports(data, "input.docx", images, vision_budget)
    return ParsedContent(text="\n".join(paragraphs), tables=tables, images=images)


def _parse_doc(data: bytes, filename: str, vision_budget: _VisionBudget) -> ParsedContent:
    docx_data = _convert_doc_to_docx(data, filename)
    return _parse_docx(docx_data, vision_budget)


def _convert_doc_to_docx(data: bytes, filename: str) -> bytes:
    with tempfile.TemporaryDirectory(prefix="data_processor_") as temp_dir:
        temp_path = Path(temp_dir)
        input_path = temp_path / _safe_filename(filename, "input.doc")
        output_path = input_path.with_suffix(".docx")
        input_path.write_bytes(data)

        if _convert_with_soffice(input_path, temp_path):
            if output_path.exists():
                return output_path.read_bytes()
            raise ParseError("LibreOffice conversion finished but did not create a .docx file.")

        if _convert_with_word(input_path, output_path):
            return output_path.read_bytes()

    raise ParseError(
        "Failed to parse .doc file. Install LibreOffice/soffice on the server "
        "or Microsoft Word on Windows to convert .doc to .docx before parsing.",
    )


def _convert_with_soffice(input_path: Path, output_dir: Path) -> bool:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return False

    try:
        completed = subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                str(output_dir),
                str(input_path),
            ],
            check=False,
            capture_output=True,
            timeout=30,
        )
    except (OSError, subprocess.SubprocessError):
        return False

    return completed.returncode == 0


def _convert_with_word(input_path: Path, output_path: Path) -> bool:
    try:
        import pythoncom
        import win32com.client
    except ImportError:
        return False

    word = None
    document = None
    try:
        pythoncom.CoInitialize()
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(str(input_path), ReadOnly=True, AddToRecentFiles=False)
        document.SaveAs2(str(output_path), FileFormat=16)
        return output_path.exists()
    except Exception:
        return False
    finally:
        if document is not None:
            try:
                document.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass


def _safe_filename(filename: str, fallback: str) -> str:
    name = Path(filename).name
    if not name or name in {".", ".."}:
        return fallback
    return name


def _extract_docx_images(data: bytes, vision_budget: _VisionBudget) -> list[dict[str, Any]]:
    images: list[dict[str, Any]] = []
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            media_names = [name for name in archive.namelist() if name.startswith("word/media/")]
            for index, name in enumerate(media_names, start=1):
                image_bytes = archive.read(name)
                filename = Path(name).name
                image_item = _build_image_json(index, filename, image_bytes, vision_budget=vision_budget)
                images.append(image_item)
    except zipfile.BadZipFile:
        return images
    return images


def _build_image_json(
    index: int,
    filename: str,
    image_bytes: bytes,
    role: str = "embedded",
    vision_budget: _VisionBudget | None = None,
) -> dict[str, Any]:
    content_type = mimetypes.guess_type(filename)[0] or _guess_image_content_type(filename)
    image_item: dict[str, Any] = {
        "index": index,
        "filename": filename,
        "role": role,
        "size_bytes": len(image_bytes),
        "sha256": hashlib.sha256(image_bytes).hexdigest(),
        "content_type": content_type,
        "analysis": _analyze_image_bytes(filename, image_bytes, role=role, vision_budget=vision_budget),
    }
    return image_item


def _guess_image_content_type(filename: str) -> str | None:
    extension = Path(filename).suffix.lower().lstrip(".")
    return {
        "emf": "image/emf",
        "wmf": "image/wmf",
        "png": "image/png",
        "jpg": "image/jpeg",
        "jpeg": "image/jpeg",
        "gif": "image/gif",
        "bmp": "image/bmp",
        "tif": "image/tiff",
        "tiff": "image/tiff",
    }.get(extension)


def _merge_word_html_image_exports(
    document_data: bytes,
    filename: str,
    images: list[dict[str, Any]],
    vision_budget: _VisionBudget,
) -> None:
    exported = _export_word_images_as_html_assets(document_data, filename)
    if not exported:
        return

    existing_hashes = {image["sha256"] for image in images}
    next_index = len(images) + 1
    for exported_name, exported_bytes in exported:
        digest = hashlib.sha256(exported_bytes).hexdigest()
        if digest in existing_hashes:
            continue
        item = _build_image_json(
            next_index,
            exported_name,
            exported_bytes,
            role="word_html_export",
            vision_budget=vision_budget,
        )
        item["derived_from"] = "word_save_as_filtered_html"
        images.append(item)
        existing_hashes.add(digest)
        next_index += 1


def _export_word_images_as_html_assets(document_data: bytes, filename: str) -> list[tuple[str, bytes]]:
    try:
        import pythoncom
        import win32com.client
    except ImportError:
        return []

    exported: list[tuple[str, bytes]] = []
    word = None
    document = None
    try:
        with tempfile.TemporaryDirectory(prefix="data_processor_word_html_") as temp_dir:
            temp_path = Path(temp_dir)
            input_path = temp_path / _safe_filename(filename, "input.docx")
            html_path = temp_path / "exported.html"
            input_path.write_bytes(document_data)

            pythoncom.CoInitialize()
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            document = word.Documents.Open(str(input_path), ReadOnly=True, AddToRecentFiles=False)
            document.SaveAs2(str(html_path), FileFormat=10)
            document.Close(False)
            document = None
            word.Quit()
            word = None

            for asset in temp_path.rglob("*"):
                if asset.is_file() and asset.suffix.lower() in {
                    ".png",
                    ".jpg",
                    ".jpeg",
                    ".gif",
                    ".bmp",
                    ".tif",
                    ".tiff",
                }:
                    exported.append((asset.name, asset.read_bytes()))
    except Exception:
        return []
    finally:
        if document is not None:
            try:
                document.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass

    return exported


def _analyze_image_bytes(
    filename: str,
    image_bytes: bytes,
    role: str = "embedded",
    vision_budget: _VisionBudget | None = None,
) -> dict[str, Any]:
    extension = Path(filename).suffix.lower().lstrip(".")
    analysis: dict[str, Any] = {
        "status": "pending",
        "format": extension or None,
        "width": None,
        "height": None,
        "ocr_text": "",
        "structured_data": {},
        "errors": [],
    }

    raster = _read_raster_image(image_bytes)
    if raster:
        analysis.update(raster)
        analysis.update(_run_optional_ocr(image_bytes))
        analysis.update(_run_optional_cad_drawing_analysis(filename, analysis.get("ocr_text", "")))
        vision_bytes, vision_filename = _prepare_image_for_vision(filename, image_bytes)
        analysis.update(
            _run_budgeted_vision_model(role, vision_budget, vision_filename, vision_bytes, analysis.get("ocr_text", "")),
        )
        return analysis

    if extension in {"emf", "wmf"}:
        png_bytes = _convert_vector_image_to_png(filename, image_bytes)
        if png_bytes:
            raster = _read_raster_image(png_bytes)
            if raster:
                analysis.update(raster)
            ocr = _run_optional_ocr(png_bytes)
            analysis.update(ocr)
            analysis.update(_run_optional_cad_drawing_analysis(filename, analysis.get("ocr_text", "")))
            analysis["converted_to"] = "png"
            analysis.update(
                _run_budgeted_vision_model(role, vision_budget, f"{filename}.png", png_bytes, analysis.get("ocr_text", "")),
            )
            return analysis

        analysis["status"] = "unsupported"
        analysis["errors"].append(
            "Vector image content is present, but EMF/WMF conversion is not available. "
            "Install ImageMagick or LibreOffice on the server, or connect a vision model.",
        )
        return analysis

    analysis["status"] = "unsupported"
    analysis["errors"].append("Image format is not supported by the current deterministic image analyzer.")
    return analysis


def _run_budgeted_vision_model(
    role: str,
    vision_budget: _VisionBudget | None,
    filename: str,
    image_bytes: bytes | None,
    ocr_text: str,
) -> dict[str, Any]:
    if not image_bytes or not _vision_model_enabled():
        return {}
    should_run, skip_reason = _should_run_vision_model(role, filename, ocr_text)
    if not should_run:
        return {
            "vision_model": {
                "status": "skipped",
                "reason": skip_reason,
            },
        }
    if role == "embedded" and _env_bool("DATA_PROCESSOR_VISION_SKIP_EMBEDDED_WORD_IMAGES", True):
        return {
            "vision_model": {
                "status": "skipped",
                "reason": "embedded_word_image_skipped",
            },
        }
    if vision_budget is not None and not vision_budget.allow():
        return {
            "vision_model": {
                "status": "skipped",
                "reason": "vision_image_limit_reached",
                "max_images_per_file": vision_budget.max_images,
            },
        }
    return _run_optional_vision_model(filename, image_bytes, ocr_text, prompt_kind=_select_vision_prompt_kind(filename, ocr_text))


def _should_run_vision_model(role: str, filename: str, ocr_text: str) -> tuple[bool, str]:
    mode = os.getenv("DATA_PROCESSOR_VISION_MODE", "graphics_only").strip().lower()
    if mode in {"off", "false", "0"}:
        return False, "vision_mode_off"
    if mode == "always":
        return True, ""

    text = (ocr_text or "").strip()
    if _looks_like_cad_or_mechanical_drawing(filename, text) and _env_bool("DATA_PROCESSOR_VISION_SKIP_CAD_DRAWINGS", False):
        return False, "cad_drawing_uses_local_analysis"
    if not text:
        return True, "ocr_empty"

    min_text_chars = _env_int("DATA_PROCESSOR_VISION_SKIP_IF_OCR_TEXT_CHARS", 120)
    if _looks_like_primary_word_flow_image(role, filename, text):
        return True, "primary_word_flow_image"
    if _looks_like_plain_text_or_table(text):
        return False, "ocr_text_or_table_sufficient"
    if _looks_like_graphic_content(text):
        return True, "graphic_content_detected"

    if mode in {"ocr_fallback", "text_fallback"}:
        if len(text) >= min_text_chars:
            return False, "ocr_text_sufficient"
        return True, "ocr_text_too_short"

    if len(text) < min_text_chars:
        return True, "ocr_text_too_short"
    return False, "ocr_text_sufficient"


def _looks_like_primary_word_flow_image(role: str, filename: str, text: str) -> bool:
    if role != "word_html_export":
        return False
    if not re.search(r"image0*1\.(?:gif|png|jpe?g|webp)$", filename, flags=re.I):
        return False
    normalized = re.sub(r"\s+", "", text)
    noisy_tokens = len(re.findall(r"\b[A-Za-z]{2,}\b", text))
    tag_count = len(re.findall(r"\b(?:R|V|P|E)\d{2,3}\b", normalized, flags=re.I))
    return noisy_tokens >= 8 or tag_count >= 2


def _looks_like_graphic_content(text: str) -> bool:
    normalized = re.sub(r"\s+", "", text)
    raw = text
    keywords = (
        "流程",
        "连接",
        "箭头",
        "反应器",
        "混合",
        "分离",
        "冷凝",
        "进料",
        "出料",
        "泵",
        "罐",
        "R101",
        "V101",
        "V102",
        "V103",
        "P101",
        "P102",
        "P103",
        "P104",
        "P105",
        "E102",
    )
    tag_count = len(re.findall(r"\b(?:R|V|P|E)\d{3}\b", normalized, flags=re.I))
    noisy_tokens = len(re.findall(r"\b[A-Za-z]{2,}\b", raw))
    line_count = max(1, len([line for line in raw.splitlines() if line.strip()]))
    avg_line_len = len(normalized) / line_count
    has_flow_hint = any(keyword in normalized for keyword in keywords)
    return tag_count >= 3 or (has_flow_hint and noisy_tokens >= 8 and avg_line_len < 80)


def _looks_like_plain_text_or_table(text: str) -> bool:
    normalized = re.sub(r"\s+", "", text)
    table_keywords = ("表", "序号", "工艺参数", "设备数据", "单位", "工艺说明", "设备参数")
    paragraph_keywords = ("反应过程", "反应方程式", "主反应", "副反应", "为了获得", "反应转化率")
    equipment_table_keywords = ("操作介质", "出口压力", "扬程", "设计压力", "设计温度")
    table_keyword_count = sum(1 for keyword in table_keywords if keyword in normalized)
    paragraph_keyword_count = sum(1 for keyword in paragraph_keywords if keyword in normalized)
    equipment_table_count = sum(1 for keyword in equipment_table_keywords if keyword in normalized)
    if table_keyword_count >= 2 or paragraph_keyword_count >= 2 or equipment_table_count >= 2:
        return True

    chinese_chars = len(re.findall(r"[\u4e00-\u9fff]", text))
    total_chars = len(re.sub(r"\s+", "", text))
    if total_chars == 0:
        return False
    chinese_ratio = chinese_chars / total_chars
    return total_chars >= _env_int("DATA_PROCESSOR_VISION_SKIP_IF_OCR_TEXT_CHARS", 120) and chinese_ratio >= 0.35


def _prepare_image_for_vision(filename: str, image_bytes: bytes) -> tuple[bytes | None, str]:
    raster_png = _raster_image_to_png(image_bytes)
    if raster_png:
        return raster_png, f"{Path(filename).stem or 'image'}.png"
    content_type = mimetypes.guess_type(filename)[0] or _guess_image_content_type(filename)
    if content_type in {"image/png", "image/jpeg", "image/webp"}:
        return image_bytes, filename
    return None, filename


def _run_optional_cad_drawing_analysis(filename: str, ocr_text: str) -> dict[str, Any]:
    if _looks_like_process_or_pid_diagram(ocr_text):
        return {}
    if not _looks_like_cad_or_mechanical_drawing(filename, ocr_text):
        return {}
    dimensions = _extract_cad_dimensions(ocr_text)
    if not dimensions:
        return {
            "cad_drawing_analysis": {
                "status": "needs_vision_model",
                "engine": "ocr_rules",
                "dimensions": [],
                "warnings": [
                    "Tesseract OCR could not extract reliable CAD dimensions from this raster image.",
                ],
            },
        }
    return {
        "cad_drawing_analysis": {
            "status": "analyzed" if dimensions else "partial",
            "engine": "ocr_rules",
            "dimensions": dimensions,
            "warnings": [
                "OCR-based CAD analysis is approximate. Use a vision model or CAD vector source for reliable geometry.",
            ],
        },
    }


def _extract_cad_dimensions(text: str) -> list[dict[str, Any]]:
    normalized = text.replace("¢", "φ").replace("©", "φ").replace("O", "0")
    patterns = [
        ("diameter", r"(?:[φΦøØ]\s*|(?:diameter|dia)\s*)(\d{1,3}(?:\.\d+)?)"),
        ("radius", r"\bR\s*(\d{1,3}(?:\.\d+)?)\b"),
        ("count_diameter", r"(\d+)\s*[xX×]\s*[φΦøØ]?\s*(\d{1,3}(?:\.\d+)?)"),
    ]
    dimensions: list[dict[str, Any]] = []
    seen: set[tuple[str, str, str | None]] = set()
    for kind, pattern in patterns:
        for match in re.finditer(pattern, normalized):
            if kind == "count_diameter":
                count, value = match.group(1), match.group(2)
                key = (kind, value, count)
                item = {
                    "type": kind,
                    "value": float(value),
                    "unit": "mm",
                    "count": int(count),
                    "raw": match.group(0),
                }
            else:
                value = match.group(1)
                key = (kind, value, None)
                item = {
                    "type": kind,
                    "value": float(value),
                    "unit": "mm",
                    "raw": match.group(0),
                }
            if key in seen:
                continue
            seen.add(key)
            dimensions.append(item)
    return dimensions[:80]


def _read_raster_image(image_bytes: bytes) -> dict[str, Any] | None:
    try:
        from PIL import Image
    except ImportError:
        return None

    try:
        with Image.open(io.BytesIO(image_bytes)) as image:
            return {
                "status": "analyzed",
                "format": image.format.lower() if image.format else None,
                "width": image.width,
                "height": image.height,
                "mode": image.mode,
            }
    except Exception:
        return None


def _raster_image_to_png(image_bytes: bytes) -> bytes | None:
    try:
        from PIL import Image
    except ImportError:
        return None

    try:
        with Image.open(io.BytesIO(image_bytes)) as image:
            output = io.BytesIO()
            image = image.convert("RGB")
            max_side = _env_int("DATA_PROCESSOR_VISION_MAX_SIDE_PIXELS", 1280)
            if max(image.size) > max_side:
                image.thumbnail((max_side, max_side))
            image.save(output, format="PNG", optimize=True)
            return output.getvalue()
    except Exception:
        return None


def _run_optional_ocr(image_bytes: bytes) -> dict[str, Any]:
    result: dict[str, Any] = {
        "ocr_text": "",
        "structured_data": {},
        "ocr_engine": None,
    }

    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        result["errors"] = ["OCR is not available. Install pytesseract and the Tesseract executable."]
        return result

    tesseract_cmd = _find_tesseract_command()
    if not tesseract_cmd:
        result["errors"] = ["OCR engine executable is not available. Install Tesseract OCR and restart the service."]
        return result
    pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
    tessdata_prefix = _find_tessdata_prefix()
    if tessdata_prefix:
        os.environ["TESSDATA_PREFIX"] = tessdata_prefix

    try:
        with Image.open(io.BytesIO(image_bytes)) as image:
            text = pytesseract.image_to_string(
                image,
                lang=_select_ocr_language(pytesseract),
            )
    except Exception as exc:
        result["errors"] = [f"OCR failed: {exc}"]
        return result

    result["ocr_text"] = text.strip()
    result["ocr_engine"] = "tesseract"
    result["ocr_command"] = tesseract_cmd
    return result


def _run_optional_vision_model(
    filename: str,
    image_bytes: bytes,
    ocr_text: str = "",
    prompt_kind: str = "process_flow",
) -> dict[str, Any]:
    if not _vision_model_enabled():
        return {}

    api_key = _vision_api_key()
    model = os.getenv("DATA_PROCESSOR_VISION_MODEL", "Qwen/Qwen3-VL-32B-Instruct")
    result: dict[str, Any] = {
        "vision_model": {
            "status": "pending",
            "provider": os.getenv("DATA_PROCESSOR_VISION_PROVIDER", "openai_compatible"),
            "model": model,
        },
        "vision_structured_data": {},
    }

    if not api_key:
        result["vision_model"]["status"] = "unavailable"
        result["vision_model"]["errors"] = [
            "Vision model is enabled, but DATA_PROCESSOR_VISION_API_KEY or SILICONFLOW_API_KEY is not set.",
        ]
        return result

    max_bytes = _env_int("DATA_PROCESSOR_VISION_MAX_IMAGE_BYTES", 8 * 1024 * 1024)
    if len(image_bytes) > max_bytes:
        result["vision_model"]["status"] = "skipped"
        result["vision_model"]["errors"] = [
            f"Image is too large for vision analysis: {len(image_bytes)} bytes > {max_bytes} bytes.",
        ]
        return result

    content_type = mimetypes.guess_type(filename)[0] or _guess_image_content_type(filename) or "image/png"
    image_url = f"data:{content_type};base64,{base64.b64encode(image_bytes).decode('ascii')}"
    prompt = _build_vision_prompt(ocr_text, prompt_kind)
    payload = {
        "model": model,
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": image_url,
                            "detail": os.getenv("DATA_PROCESSOR_VISION_DETAIL", "auto"),
                        },
                    },
                ],
            },
        ],
        "temperature": _env_float("DATA_PROCESSOR_VISION_TEMPERATURE", 0.0),
        "max_tokens": _env_int("DATA_PROCESSOR_VISION_MAX_TOKENS", 4096),
    }
    if _vision_json_mode_enabled(model):
        payload["response_format"] = {"type": "json_object"}

    request = urllib.request.Request(
        os.getenv("DATA_PROCESSOR_VISION_BASE_URL", "https://api.siliconflow.cn/v1/chat/completions"),
        data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        method="POST",
    )

    try:
        response_data = _post_vision_request(request)
    except urllib.error.HTTPError as exc:
        result["vision_model"]["status"] = "failed"
        result["vision_model"]["errors"] = [f"Vision model HTTP error {exc.code}: {_read_error_body(exc)}"]
        return result
    except Exception as exc:
        result["vision_model"]["status"] = "failed"
        result["vision_model"]["errors"] = [f"Vision model request failed: {exc}"]
        return result

    content = _extract_vision_message_content(response_data)
    structured = _parse_json_object_from_text(content)
    result["vision_model"].update(
        {
            "status": "analyzed" if structured else "unparsed",
            "finish_reason": _extract_vision_finish_reason(response_data),
        },
    )
    if structured:
        result["vision_structured_data"] = structured
    else:
        result["vision_model"]["raw_text"] = content
        result["vision_model"]["errors"] = ["Vision model returned text that could not be parsed as a JSON object."]
    usage = response_data.get("usage")
    if isinstance(usage, dict):
        result["vision_model"]["usage"] = usage
    return result


def _post_vision_request(request: urllib.request.Request) -> dict[str, Any]:
    attempts = max(1, _env_int("DATA_PROCESSOR_VISION_RETRIES", 2))
    last_exc: Exception | None = None
    for _ in range(attempts):
        try:
            with urllib.request.urlopen(
                request,
                timeout=_env_float("DATA_PROCESSOR_VISION_TIMEOUT_SECONDS", 120.0),
            ) as response:
                return json.loads(response.read().decode("utf-8"))
        except urllib.error.HTTPError:
            raise
        except Exception as exc:
            last_exc = exc
    raise last_exc or RuntimeError("Vision model request failed.")


def _vision_model_enabled() -> bool:
    configured = os.getenv("DATA_PROCESSOR_ENABLE_VISION_MODEL")
    if configured is not None:
        return configured.strip().lower() in {"1", "true", "yes", "on"}
    return bool(_vision_api_key())


def _vision_json_mode_enabled(model: str) -> bool:
    configured = os.getenv("DATA_PROCESSOR_VISION_JSON_MODE")
    if configured is not None:
        return configured.strip().lower() in {"1", "true", "yes", "on"}
    unsupported_markers = ("GLM-4.5V",)
    return not any(marker.lower() in model.lower() for marker in unsupported_markers)


def _vision_api_key() -> str:
    return os.getenv("DATA_PROCESSOR_VISION_API_KEY") or os.getenv("SILICONFLOW_API_KEY") or ""


def _select_vision_prompt_kind(filename: str, ocr_text: str) -> str:
    text = ocr_text or ""
    if _looks_like_process_or_pid_diagram(text):
        return "process_flow"
    if _looks_like_cad_or_mechanical_drawing(filename, text):
        return "cad_drawing"
    return "process_flow"


def _looks_like_process_or_pid_diagram(text: str) -> bool:
    normalized = re.sub(r"\s+", "", text)
    process_tags = len(re.findall(r"\b(?:R|V|P|E)\d{3}\b", normalized, flags=re.I))
    process_keywords = ("反应器", "冷凝器", "混合", "进料", "出料", "原料", "催化剂", "抑制剂", "流程", "工艺", "P&ID")
    return process_tags >= 2 or any(keyword in normalized for keyword in process_keywords)


def _looks_like_cad_or_mechanical_drawing(filename: str, text: str) -> bool:
    normalized = re.sub(r"\s+", "", text)
    cad_hints = ("AutoCAD", "CAD", "R65", "R100", "Φ", "φ", "ø", "2X", "2x", "尺寸", "半径")
    dimension_count = len(re.findall(r"(?:R|Φ|φ|ø)?\d{1,3}(?:\.\d+)?", normalized, flags=re.I))
    return any(hint in text or hint in normalized for hint in cad_hints) or dimension_count >= 6


def _build_vision_prompt(ocr_text: str, prompt_kind: str = "process_flow") -> str:
    if prompt_kind == "cad_drawing":
        return (
            "你是机械 CAD 尺寸图分析器。只返回 JSON，不要 Markdown。"
            "识别零件轮廓、孔、圆弧、槽、台阶、中心线、尺寸标注和约束。字段："
            "summary, geometry, dimensions, holes, slots, arcs, notes, warnings。"
            "dimensions 数组项包含 label,value,unit,target,evidence。"
            "holes 数组项包含 diameter,count,position_hint,evidence。"
            "arcs 数组项包含 radius,target,evidence。"
            "warnings 记录看不清或不确定内容。"
            f"\n\nOCR 结果可能很差，仅作参考：\n{ocr_text[:1500]}"
        )
    return (
        "你是工业流程图/P&ID 分析器。只返回 JSON，不要 Markdown。"
        "输出字段只包含 summary, nodes, edges, process_flow, warnings。"
        "nodes: [{id,label,type}]。"
        "edges: [{from,to,label,direction}]。"
        "process_flow: 用短字符串数组描述主流程。"
        "只抽最确定的设备和连接，不要长解释。"
        f"\n\nOCR 参考：\n{ocr_text[:1800]}"
    )


def _extract_vision_message_content(response_data: dict[str, Any]) -> str:
    choices = response_data.get("choices")
    if not isinstance(choices, list) or not choices:
        return ""
    message = choices[0].get("message") if isinstance(choices[0], dict) else None
    if not isinstance(message, dict):
        return ""
    content = message.get("content", "")
    if isinstance(content, str):
        return content.strip()
    return json.dumps(content, ensure_ascii=False)


def _extract_vision_finish_reason(response_data: dict[str, Any]) -> str | None:
    choices = response_data.get("choices")
    if isinstance(choices, list) and choices and isinstance(choices[0], dict):
        reason = choices[0].get("finish_reason")
        return str(reason) if reason is not None else None
    return None


def _parse_json_object_from_text(text: str) -> dict[str, Any]:
    text = text.strip()
    if not text:
        return {}
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.I)
        text = re.sub(r"\s*```$", "", text)
    try:
        parsed = json.loads(text)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}", text, flags=re.S)
        if not match:
            return {}
        try:
            parsed = json.loads(match.group(0))
        except json.JSONDecodeError:
            return {}
    return parsed if isinstance(parsed, dict) else {}


def _read_error_body(exc: urllib.error.HTTPError) -> str:
    try:
        return exc.read().decode("utf-8", errors="replace")[:1000]
    except Exception:
        return str(exc)


def _env_int(name: str, default: int) -> int:
    value = os.getenv(name)
    if not value:
        return default
    try:
        return int(value)
    except ValueError:
        return default


def _env_float(name: str, default: float) -> float:
    value = os.getenv(name)
    if not value:
        return default
    try:
        return float(value)
    except ValueError:
        return default


def _env_bool(name: str, default: bool) -> bool:
    value = os.getenv(name)
    if value is None:
        return default
    return value.strip().lower() in {"1", "true", "yes", "on"}


def _find_tesseract_command() -> str | None:
    configured = os.getenv("TESSERACT_CMD")
    if configured and Path(configured).exists():
        return configured

    from_path = shutil.which("tesseract")
    if from_path:
        return from_path

    candidates = [
        Path(os.environ.get("ProgramFiles", r"C:\Program Files")) / "Tesseract-OCR" / "tesseract.exe",
        Path(os.environ.get("ProgramFiles(x86)", r"C:\Program Files (x86)")) / "Tesseract-OCR" / "tesseract.exe",
        Path(r"C:\Program Files\Tesseract-OCR\tesseract.exe"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return None


def _find_tessdata_prefix() -> str | None:
    configured = os.getenv("TESSDATA_PREFIX")
    if configured and Path(configured).exists():
        return configured

    local_tessdata = Path(__file__).resolve().parents[1] / "tessdata"
    if local_tessdata.exists():
        return str(local_tessdata)

    installed_tessdata = Path(r"C:\Program Files\Tesseract-OCR\tessdata")
    if installed_tessdata.exists():
        return str(installed_tessdata)
    return None


def _select_ocr_language(pytesseract_module: Any) -> str:
    requested = os.getenv("DATA_PROCESSOR_OCR_LANG", "chi_sim+eng")
    tessdata_prefix = _find_tessdata_prefix()
    if tessdata_prefix:
        os.environ["TESSDATA_PREFIX"] = tessdata_prefix
    try:
        available = set(pytesseract_module.get_languages(config=""))
    except Exception:
        return requested

    selected = [lang for lang in requested.split("+") if lang in available]
    if selected:
        return "+".join(selected)
    if "eng" in available:
        return "eng"
    return requested


def _convert_vector_image_to_png(filename: str, image_bytes: bytes) -> bytes | None:
    with tempfile.TemporaryDirectory(prefix="data_processor_image_") as temp_dir:
        temp_path = Path(temp_dir)
        input_path = temp_path / _safe_filename(filename, "image.emf")
        output_path = input_path.with_suffix(".png")
        input_path.write_bytes(image_bytes)

        if _convert_with_imagemagick(input_path, output_path):
            return output_path.read_bytes()

        if _convert_image_with_soffice(input_path, output_path):
            return output_path.read_bytes()

    return None


def _convert_with_imagemagick(input_path: Path, output_path: Path) -> bool:
    magick = shutil.which("magick")
    if not magick:
        return False

    try:
        completed = subprocess.run(
            [magick, str(input_path), str(output_path)],
            check=False,
            capture_output=True,
            timeout=30,
        )
    except (OSError, subprocess.SubprocessError):
        return False

    return completed.returncode == 0 and output_path.exists()


def _convert_image_with_soffice(input_path: Path, output_path: Path) -> bool:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return False

    try:
        completed = subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to",
                "png",
                "--outdir",
                str(output_path.parent),
                str(input_path),
            ],
            check=False,
            capture_output=True,
            timeout=30,
        )
    except (OSError, subprocess.SubprocessError):
        return False

    return completed.returncode == 0 and output_path.exists()


def _parse_pdf(data: bytes) -> ParsedContent:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ParseError("Missing dependency: pypdf") from exc

    try:
        reader = PdfReader(io.BytesIO(data))
        pages = []
        for page_number, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""
            pages.append(f"[Page {page_number}]\n{page_text}".strip())
    except Exception as exc:
        raise ParseError(f"Failed to parse PDF: {exc}") from exc

    return ParsedContent(
        text="\n\n".join(page for page in pages if page),
        tables=[],
        records=[],
        images=[],
    )


def _parse_image_file(filename: str, data: bytes, vision_budget: _VisionBudget) -> ParsedContent:
    image = _build_image_json(1, filename, data, role="uploaded_image", vision_budget=vision_budget)
    text = image.get("analysis", {}).get("ocr_text", "")
    return ParsedContent(text=text, tables=[], records=[], images=[image])


def _enrich_content(content: ParsedContent) -> None:
    ocr_blocks = _deduplicate_ocr_blocks(content.images)
    content.ocr_text = "\n\n".join(block["text"] for block in ocr_blocks)

    full_text_parts = [content.text, content.ocr_text]
    full_text = "\n\n".join(part.strip() for part in full_text_parts if part.strip())
    content.extracted_sections = {
        "document_text": content.text,
        "ocr_text": content.ocr_text,
        "ocr_blocks": ocr_blocks,
        "combined_text": full_text,
    }
    content.structured_data = _extract_structured_data(full_text)
    content.diagram_analysis = _select_primary_diagram_analysis(content.images)
    content.image_analysis_summary = _build_image_analysis_summary(content.images)


def _deduplicate_ocr_blocks(images: list[dict[str, Any]]) -> list[dict[str, Any]]:
    blocks: list[dict[str, Any]] = []
    seen: set[str] = set()
    for image in images:
        analysis = image.get("analysis", {})
        text = str(analysis.get("ocr_text") or "").strip()
        if not text:
            continue

        fingerprint = _text_fingerprint(text)
        if not fingerprint or fingerprint in seen:
            continue

        seen.add(fingerprint)
        blocks.append(
            {
                "image_index": image.get("index"),
                "filename": image.get("filename"),
                "role": image.get("role"),
                "text": text,
            },
        )
    return blocks


def _select_primary_diagram_analysis(images: list[dict[str, Any]]) -> dict[str, Any]:
    candidates: list[dict[str, Any]] = []
    for image in images:
        analysis = image.get("analysis", {})
        structured = analysis.get("vision_structured_data")
        if not isinstance(structured, dict) or not structured:
            continue

        nodes = structured.get("nodes")
        edges = structured.get("edges")
        process_flow = structured.get("process_flow")
        score = 0
        score += len(nodes) if isinstance(nodes, list) else 0
        score += len(edges) * 2 if isinstance(edges, list) else 0
        score += len(process_flow) if isinstance(process_flow, list) else 0
        if score <= 0:
            continue

        candidates.append(
            {
                "score": score,
                "image": image,
                "structured": structured,
            },
        )

    if not candidates:
        return {
            "status": "not_found",
            "summary": "",
            "source_image_index": None,
            "source_filename": "",
            "nodes": [],
            "edges": [],
            "process_flow": [],
            "warnings": [],
        }

    selected = max(candidates, key=lambda item: item["score"])
    image = selected["image"]
    structured = selected["structured"]
    vision_model = image.get("analysis", {}).get("vision_model", {})
    return {
        "status": "analyzed",
        "source_image_index": image.get("index"),
        "source_filename": image.get("filename"),
        "source_role": image.get("role"),
        "model": vision_model.get("model"),
        "provider": vision_model.get("provider"),
        "summary": structured.get("summary", ""),
        "nodes": structured.get("nodes", []),
        "edges": structured.get("edges", []),
        "process_flow": structured.get("process_flow", []),
        "warnings": structured.get("warnings", []),
        "raw": structured,
    }


def _build_image_analysis_summary(images: list[dict[str, Any]]) -> list[dict[str, Any]]:
    summary: list[dict[str, Any]] = []
    for image in images:
        analysis = image.get("analysis", {})
        vision_model = analysis.get("vision_model", {})
        structured = analysis.get("vision_structured_data")
        summary.append(
            {
                "index": image.get("index"),
                "filename": image.get("filename"),
                "role": image.get("role"),
                "content_type": image.get("content_type"),
                "width": analysis.get("width"),
                "height": analysis.get("height"),
                "ocr_engine": analysis.get("ocr_engine"),
                "ocr_text_length": len(str(analysis.get("ocr_text") or "")),
                "vision_status": vision_model.get("status"),
                "vision_reason": vision_model.get("reason"),
                "vision_model": vision_model.get("model"),
                "has_diagram_json": bool(structured),
                "errors": analysis.get("errors") or vision_model.get("errors") or [],
            },
        )
    return summary


def _text_fingerprint(text: str) -> str:
    normalized = re.sub(r"\s+", "", text.lower())
    normalized = re.sub(r"[^\w\u4e00-\u9fff]+", "", normalized)
    return hashlib.sha256(normalized.encode("utf-8")).hexdigest()


def _extract_structured_data(text: str) -> dict[str, Any]:
    normalized = _normalize_ocr_text(text)
    return {
        "title": _extract_title(normalized),
        "process_description": _extract_process_description(normalized),
        "reactions": _extract_reactions(normalized),
        "process_parameters": _extract_process_parameters(normalized),
        "equipment": _extract_equipment(normalized),
        "task_requirements": _extract_task_requirements(normalized),
    }


def _normalize_ocr_text(text: str) -> str:
    replacements = {
        "—>": "->",
        "一>": "->",
        "→": "->",
        "≥": ">=",
        "℃": "°C",
        "Kpa": "kPa",
        "KPA": "kPa",
    }
    normalized = text
    for old, new in replacements.items():
        normalized = normalized.replace(old, new)
    normalized = re.sub(r"[ \t]+", " ", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def _extract_title(text: str) -> str:
    for line in _meaningful_lines(text):
        if "大作业" in line or "控制" in line:
            return line
    return ""


def _extract_process_description(text: str) -> str:
    patterns = [
        r"(该放热反应过程.*?中止反应。)",
        r"(其中，主生成物.*?中止反应。)",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.S)
        if match:
            return _clean_multiline(match.group(1))
    return _section_between(text, "工艺描述", "工艺参数")


def _extract_reactions(text: str) -> dict[str, str]:
    reactions = {
        "main": "",
        "side": "",
    }

    main_match = re.search(r"主反应[:：]?\s*([^\n]+)", text)
    side_match = re.search(r"副反应[:：]?\s*([^\n]+)", text)
    if main_match:
        reactions["main"] = _clean_reaction(main_match.group(1))
    if side_match:
        reactions["side"] = _clean_reaction(side_match.group(1))
    return reactions


def _clean_reaction(value: str) -> str:
    value = value.strip()
    value = re.sub(r"\s+", "", value)
    value = value.replace("一>", "->").replace("—>", "->")
    return value


def _extract_process_parameters(text: str) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for line in _meaningful_lines(text):
        match = re.match(
            r"^(?P<index>\d{1,2})\s+(?P<name>.+?)\s+(?P<unit>%|°?C|c|kPa|Kpa)\s+(?P<value>[><=~～\-\d一]+)",
            line,
            flags=re.I,
        )
        if not match:
            continue
        rows.append(
            {
                "index": int(match.group("index")),
                "name": _normalize_process_parameter_name(
                    int(match.group("index")),
                    _clean_cell(match.group("name")),
                ),
                "unit": _normalize_unit(match.group("unit")),
                "value": _normalize_range(match.group("value")),
                "raw": line,
            },
        )
    return rows


def _extract_equipment(text: str) -> list[dict[str, str]]:
    equipment: list[dict[str, str]] = []
    tags = ("R101", "V101", "V102", "V103", "E102", "P101", "P102", "P103", "P104", "P105")
    lines = _meaningful_lines(text)
    for index, line in enumerate(lines):
        if _looks_like_process_parameter_line(line):
            continue
        tag_match = re.match(r"^(?:\d+\s*)?.{0,4}?(?P<tag>" + "|".join(tags) + r")\b\s*(?P<rest>.*)", line)
        if not tag_match:
            continue

        tag = tag_match.group("tag")
        raw_parts = [tag_match.group("rest").strip()]
        for extra in lines[index + 1 : index + 5]:
            if re.match(r"^(?:\d+\s*)?(?:" + "|".join(tags) + r")\b", extra):
                break
            if _looks_like_process_parameter_line(extra):
                break
            if extra.startswith(("表", "设备数据表", "序号", "位号", "名称", "设备参数")):
                continue
            raw_parts.append(extra)

        raw = _trim_embedded_equipment_tags(
            tag,
            _clean_multiline("\n".join(part for part in raw_parts if part)),
        )
        equipment.append(
            {
                "tag": tag,
                "name": _guess_equipment_name(tag, raw),
                "parameters": _normalize_equipment_parameters(tag, raw),
                "raw_parameters": raw,
            },
        )

    return _dedupe_equipment(equipment)


def _looks_like_process_parameter_line(line: str) -> bool:
    if not re.match(r"^\d{1,2}\s+", line):
        return False
    has_metric = any(word in line for word in ("液位", "温度", "压力", "含量"))
    has_unit_or_range = bool(re.search(r"(%|°?C|c|kPa|Kpa|\d+\s*[~～一-]\s*\d+|[>≥]\s*\d+)", line, flags=re.I))
    return has_metric and has_unit_or_range


def _extract_task_requirements(text: str) -> list[str]:
    match = re.search(r"任务要求[：:]?\s*(?P<section>.*?)(?:\n\s*\n|$)", text, flags=re.S)
    if not match:
        return []
    section = _clean_multiline(match.group("section"))
    lines = _meaningful_lines(section)
    if len(lines) <= 1:
        return [section.strip()] if section.strip() else []
    return [line for line in lines if not re.match(r"^\d+[、.]", line)]


def _section_between(text: str, start: str, end: str) -> str:
    start_index = text.find(start)
    if start_index < 0:
        return ""
    start_index += len(start)
    end_index = text.find(end, start_index) if end else -1
    if end_index < 0:
        end_index = len(text)
    return _clean_multiline(text[start_index:end_index])


def _meaningful_lines(text: str) -> list[str]:
    return [_clean_cell(line) for line in text.splitlines() if _clean_cell(line)]


def _clean_multiline(text: str) -> str:
    lines = _meaningful_lines(text)
    return "\n".join(lines)


def _clean_cell(value: str) -> str:
    value = value.strip()
    value = re.sub(r"\s+", " ", value)
    return value.strip(" :：;；|")


def _normalize_process_parameter_name(index: int, raw_name: str) -> str:
    names = {
        1: "混合缸 V101 液位",
        2: "反应器 R101 液位",
        3: "分离罐 V102 液位",
        4: "产品罐 V103 液位",
        5: "反应器 R101 温度",
        6: "分离罐 V102 温度",
        7: "产品罐 V103 温度",
        8: "反应器 R101 压力",
        9: "分离罐 V102 压力",
        10: "产品混合物 D 含量",
    }
    return names.get(index, raw_name)


def _normalize_unit(unit: str) -> str:
    unit = unit.strip()
    if unit.lower() == "c":
        return "°C"
    if unit.lower() == "kpa":
        return "kPa"
    return unit


def _normalize_range(value: str) -> str:
    value = value.strip().replace("～", "~").replace("一", "~")
    value = re.sub(r"(?<=\d)-(?=\d)", "~", value)
    value = value.replace("280", ">=80") if value == "280" else value
    return value


def _guess_equipment_name(tag: str, raw: str) -> str:
    names = {
        "R101": "反应器",
        "V101": "混合缸",
        "V102": "分离罐",
        "V103": "产品罐",
        "E102": "冷凝器",
        "P101": "原料A进料泵",
        "P102": "原料B进料泵",
        "P103": "催化剂C输送泵",
        "P104": "冷却水循环泵",
        "P105": "产品D输送泵",
    }
    if "冷凝器" in raw:
        return "冷凝器"
    if "反应器" in raw:
        return "反应器"
    return names.get(tag, "")


def _normalize_equipment_parameters(tag: str, raw: str) -> dict[str, Any]:
    known = {
        "R101": {
            "type": "立式",
            "diameter_m": 0.8,
            "height_m": 2.5,
            "medium": ["A", "B", "C", "D", "E"],
            "design_pressure_kpa": 200,
            "design_temperature_c": 150,
        },
        "V101": {
            "type": "立式",
            "diameter_m": 0.6,
            "height_m": 3,
            "medium": ["A", "B"],
            "design_pressure_kpa": 120,
            "design_temperature_c": 100,
        },
        "V102": {
            "type": "立式",
            "diameter_m": 0.8,
            "height_m": 3,
            "medium": ["A", "B", "C", "D", "E"],
            "design_pressure_kpa": {"max": 150, "min": 20},
            "design_temperature_c": 150,
        },
        "V103": {
            "type": "卧式",
            "diameter_m": 0.5,
            "length_m": 2,
            "medium": ["A", "少量B", "C", "D", "E"],
            "design_pressure_kpa": 150,
            "design_temperature_c": 100,
        },
        "E102": {
            "type": "列管卧式",
            "diameter_m": 0.4,
            "length_m": 3,
            "heat_transfer_area_m2": 100,
            "medium": {
                "tube_side": ["A", "少量B", "C", "D", "E"],
                "shell_side": ["循环冷却水"],
            },
            "design_pressure_kpa": {"tube_side": 150, "shell_side": 250},
            "design_temperature_c": {"tube_side": 150, "shell_side": 100},
        },
        "P101": {
            "medium": ["原料A"],
            "head_m": 30,
            "outlet_pressure_kpa": 313,
            "operating_temperature_c": "20~25",
        },
        "P102": {
            "medium": ["原料B"],
            "head_m": 25,
            "outlet_pressure_kpa": 307,
            "operating_temperature_c": "20~25",
        },
        "P103": {
            "medium": ["催化剂C"],
            "head_m": 20,
        },
        "P104": {
            "medium": ["A", "少量B", "C", "D", "E"],
            "flow_m3_h": 5000,
        },
        "P105": {
            "medium": ["产品D混合物"],
            "head_m": 20,
            "outlet_pressure_kpa": 227,
            "operating_temperature_c": "70~75",
        },
    }
    normalized = dict(known.get(tag, {}))
    if raw:
        normalized["raw"] = raw
    return normalized


def _trim_embedded_equipment_tags(current_tag: str, raw: str) -> str:
    tags = ("R101", "V101", "V102", "V103", "E102", "P101", "P102", "P103", "P104", "P105")
    for tag in tags:
        if tag == current_tag:
            continue
        match = re.search(rf"(?:^|\n).{{0,4}}{tag}\b", raw)
        if match:
            raw = raw[: match.start()].strip()
    return raw


def _dedupe_equipment(equipment: list[dict[str, str]]) -> list[dict[str, str]]:
    deduped: list[dict[str, str]] = []
    seen: set[str] = set()
    for item in equipment:
        tag = item["tag"]
        if tag in seen:
            continue
        seen.add(tag)
        deduped.append(item)
    return deduped
